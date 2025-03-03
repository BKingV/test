import streamlit as st
from docx import Document

st.title("📄 Онлайн-тестирование по темам")

uploaded_file = st.file_uploader("Загрузите Word-файл с тестами", type=["docx"])

def extract_themes_and_questions(doc):
    """Извлекает темы, подтемы и вопросы, начиная обработку только с первой темы, после которой идет таблица"""
    themes = {}
    
    st.write("📌 Количество таблиц в документе:", len(doc.tables))  # Проверяем количество таблиц
    tables_iter = iter(doc.tables)

    for para in doc.paragraphs:
        text = para.text.strip()

        if text.startswith("ТЕМА:"):  
            current_theme = text.replace("ТЕМА:", "").strip()
            themes[current_theme] = {}

            try:
                table = next(tables_iter)  
                rows = table.rows
                if len(rows) < 2:
                    continue  

                headers = [cell.text.strip().lower() for cell in rows[0].cells]
                if "текст вопроса" not in headers or "варианты ответа" not in headers:
                    continue  

                question_idx = headers.index("текст вопроса")
                answers_idx = headers.index("варианты ответа")
                correct_idx = headers.index("эталон") if "эталон" in headers else None

                current_subtheme = None

                for row in rows[1:]:
                    first_cell_text = row.cells[0].text.strip()
                    question_text = row.cells[question_idx].text.strip()
                    answer_text = row.cells[answers_idx].text.strip()
                    correct_text = row.cells[correct_idx].text.strip() if correct_idx else ""

                    # Если строка содержит только название подтемы, сохраняем ее
                    if first_cell_text and not question_text:
                        current_subtheme = first_cell_text
                        themes[current_theme][current_subtheme] = []
                        continue

                    # Если есть вопрос, добавляем его к текущей подтеме (или теме)
                    if current_subtheme:
                        target_list = themes[current_theme][current_subtheme]
                    else:
                        target_list = themes[current_theme].setdefault("Без подтем", [])

                    question_entry = {
                        "question": question_text,
                        "answers": [],
                        "correct": []
                    }

                    if not target_list or target_list[-1]["question"] != question_text:
                        target_list.append(question_entry)

                    target_list[-1]["answers"].append(answer_text)
                    if correct_text:
                        target_list[-1]["correct"].append(answer_text)

            except StopIteration:
                pass  

    return themes

if uploaded_file:
    doc = Document(uploaded_file)
    themes = extract_themes_and_questions(doc)

    if not themes:
        st.warning("⚠️ Не удалось извлечь темы и вопросы. Проверьте формат документа.")
    else:
        if "themes" not in st.session_state:
            st.session_state["themes"] = themes
            st.session_state["selected_theme"] = None
            st.session_state["selected_subtheme"] = None
            st.session_state["questions"] = []
            st.session_state["current_question"] = 0
            st.session_state["test_started"] = False
            st.session_state["show_result"] = False
            st.session_state["selected_answers"] = {}
            st.session_state["show_confirm_exit"] = False

        st.subheader("📚 Выберите тему:")
        selected_theme = st.selectbox("Выберите тему", list(themes.keys()), key="theme_select")

        if selected_theme:
            subthemes = list(themes[selected_theme].keys())

            if subthemes and subthemes != ["Без подтем"]:
                st.subheader("📂 Выберите подтему:")
                selected_subtheme = st.selectbox("Выберите подтему", subthemes, key="subtheme_select")
            else:
                selected_subtheme = None  

            if st.button("▶️ Начать тест"):
                st.session_state["selected_theme"] = selected_theme
                st.session_state["selected_subtheme"] = selected_subtheme
                st.session_state["questions"] = themes[selected_theme].get(selected_subtheme, [])
                st.session_state["test_started"] = True
                st.session_state["current_question"] = 0
                st.session_state["selected_answers"] = {}
                st.rerun()

if st.session_state.get("test_started", False):
    q_idx = st.session_state["current_question"]
    questions = st.session_state["questions"]

    if q_idx < len(questions):
        question_data = questions[q_idx]

        st.subheader(f"Вопрос {q_idx + 1} из {len(questions)}")
        st.write(question_data["question"])

        selected_answers = st.session_state["selected_answers"].get(q_idx, [])

        for i, answer in enumerate(question_data["answers"]):
            key = f"q{q_idx}_a{i}"
            checked = answer in selected_answers
            if st.checkbox(answer, key=key, value=checked):
                if answer not in selected_answers:
                    selected_answers.append(answer)
            else:
                if answer in selected_answers:
                    selected_answers.remove(answer)

        st.session_state["selected_answers"][q_idx] = selected_answers

        col1, col2, col3 = st.columns([1, 2, 1])

        with col1:
            if q_idx > 0:
                if st.button("⬅️ Предыдущий вопрос"):
                    st.session_state["current_question"] -= 1
                    st.rerun()

        with col3:
            if q_idx + 1 < len(questions):
                if st.button("➡️ Следующий вопрос"):
                    st.session_state["current_question"] += 1
                    st.rerun()
            else:
                if st.button("✅ Завершить тест"):
                    st.session_state["show_result"] = True
                    st.rerun()

if st.session_state.get("show_result", False):
    st.subheader("📊 Результаты теста")

    results = []
    correct_count = 0
    total_questions = len(st.session_state["questions"])

    for i, question in enumerate(st.session_state["questions"]):
        user_answers = st.session_state["selected_answers"].get(i, [])
        correct_answers = question["correct"]
        is_correct = set(user_answers) == set(correct_answers)

        if is_correct:
            correct_count += 1
        else:
            results.append({
                "Вопрос": question["question"],
                "Ваш ответ": ", ".join(user_answers),
                "Правильный ответ": ", ".join(correct_answers)
            })

    st.write(f"✅ Вы ответили правильно на {correct_count} из {total_questions} вопросов.")

    if results:
        st.write("❌ Ошибки:")
        for res in results:
            with st.expander(res["Вопрос"]):
                st.write(f"**Ваш ответ:** {res['Ваш ответ']}")
                st.write(f"✅ **Правильный ответ:** {res['Правильный ответ']}")

    if st.button("🔄 Пройти еще раз"):
        st.session_state["test_started"] = False
        st.session_state["show_result"] = False
        st.session_state["selected_theme"] = None
        st.session_state["selected_subtheme"] = None
        st.session_state["questions"] = []
        st.session_state["current_question"] = 0
        st.session_state["selected_answers"] = {}
        st.rerun()
